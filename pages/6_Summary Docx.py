import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from datetime import date
from pathlib import Path
import json
from utils import secure_filename 

# --- Session Check ---
if "filename" not in st.session_state:
    st.warning("No file selected. Please go to the main page.")
    st.stop()

filename = secure_filename(st.session_state["filename"])  # ✅ sanitize
directory = Path(f"data/{filename}")
json_path = directory / f"Summary_{filename}.json"
docx_filename = directory / f"Summary_{filename}.docx"
directory.mkdir(parents=True, exist_ok=True)

# --- Load local data ---
if json_path.exists():
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        st.session_state["summary_data"] = data
    except Exception as e:
        st.error(f"❌ Could not load summary data: {e}")
        data = {}
else:
    st.warning("⚠️ Local summary file not found.")
    data = {}

st.title(f"Summary Document for {filename}")

# --- Word Styling Helper ---
def create_shading_element(color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    return shading

def create_word_doc(filename, data, output_path):
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    document.core_properties.author = "Dr. St^2"

    # Header
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filename_run = title_paragraph.add_run(filename)
    filename_run.font.name = "Arial"
    filename_run.bold = True
    filename_run.font.size = Pt(16)

    title_paragraph.add_run("\t" * 7)
    date_run = title_paragraph.add_run(f"{date.today()}")
    date_run.font.name = "Arial"
    date_run.bold = True
    date_run.font.size = Pt(16)

    # Table of key fields
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    labels = [
        "Independent Claims", "Ptbs", "Solution", "Technical Effect",
        "Keywords", "Classes", "Remarks", "Unity", "Prior Art"
    ]

    for i, label in enumerate(labels):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(data.get(label, ""))

        run_left = row[0].paragraphs[0].runs[0]
        run_left.font.name = "Arial"
        run_left.font.bold = True
        run_left.font.size = Pt(14)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        row[0].add_paragraph()

        run_right = row[1].paragraphs[0].runs[0]
        run_right.font.name = "Arial"
        run_right.font.size = Pt(12)

        color = "D9EAF7" if i % 2 == 0 else "FFFFFF"
        for cell in row:
            cell._element.get_or_add_tcPr().append(create_shading_element(color))

    # Add Image
    image_path = directory / f"appl_image_{filename}.png"
    image_row = table.add_row().cells
    image_cell = image_row[0]
    image_cell.merge(image_row[1])

    if image_path.is_file():
        with Image.open(image_path) as img:
            img_width, img_height = img.size
            max_width_mm, max_height_mm = 140, 100
            mm_to_px = lambda mm: int((mm / 25.4) * 96)
            max_width_px = mm_to_px(max_width_mm)
            max_height_px = mm_to_px(max_height_mm)
            aspect_ratio = img_width / img_height

            if img_width > max_width_px or img_height > max_height_px:
                if img_width / max_width_px > img_height / max_height_px:
                    new_width = max_width_px
                    new_height = int(new_width / aspect_ratio)
                else:
                    new_height = max_height_px
                    new_width = int(new_height * aspect_ratio)
            else:
                new_width, new_height = img_width, img_height

            new_width_mm = (new_width / 96) * 25.4
            new_height_mm = (new_height / 96) * 25.4

        para = image_cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(str(image_path), width=Mm(new_width_mm), height=Mm(new_height_mm))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        image_cell.text = "You did not provide an application image."
        run = image_cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0)
        image_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(output_path)

# --- Display RoSS (Summary View) ---
ross_data = {key: data.get(key, "") for key in [
    "Ptbs", "Technical Effect", "Solution", "Keywords", "Classes", "Remarks"
]}

combinations = data.get("Markers", {}).get("Combinations", [])
ross_data["Markers"] = ", ".join(combinations) if isinstance(combinations, list) else ""

ross_text = "\n".join(str(v) for v in ross_data.values()).strip()

st.text_area(
    "RoSS Summary",
    value=ross_text,
    height=200,
    placeholder="If this is empty, add info in the 'General' tab"
)

# --- Create + Download DOCX ---
if st.button("📄 Create and Download Word", type="primary", use_container_width=True):
    try:
        create_word_doc(filename, data, docx_filename)
        with open(docx_filename, "rb") as f:
            word_bytes = f.read()

        st.download_button(
            label="⬇️ Download Word Document",
            data=word_bytes,
            file_name=f"Summary_{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("✅ Word document created and ready for download.")
    except Exception as e:
        st.error(f"❌ Error creating or loading Word document: {e}")
